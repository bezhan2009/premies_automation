import os
import time
import zipfile
from datetime import datetime, date
from typing import List

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from configs.load_configs import get_config
from internal.app.models.application import ApplicationInfo
from internal.lib.date import get_current_date
from internal.lib.file import sanitize_filename
from internal.repository.application import get_application_by_ids
from internal.service.automation.base_automation import BaseAutomation
from pkg.errors.not_found_error import NotFoundError
from pkg.logger.logger import setup_logger

logger = setup_logger(__name__)


class PollAutomation(BaseAutomation):
    def __init__(self):
        self.automation_configs = get_config().automation_details
        self.file_path = ""
        self.template_path = self.automation_configs.def_template_paths.def_report_template_poll
        super().__init__("service.automation.PollAutomation")

    def _get_application_data(self, application_ids: List[int]) -> List[ApplicationInfo]:
        return get_application_by_ids(application_ids)

    def _format_date(self, value):
        if isinstance(value, datetime):
            return value.replace(tzinfo=None).strftime("%d.%m.%Y")
        elif isinstance(value, date):
            return value.strftime("%d.%m.%Y")
        return value or ""

    def _format_address(self, app: ApplicationInfo) -> str:
        address_parts = []
        if app.country:
            address_parts.append(app.country)
        if app.region:
            address_parts.append(app.region)
        if app.district:
            address_parts.append(app.district)
        if app.populated:
            address_parts.append(f"{app.population_type} {app.populated}")
        if app.street:
            address_parts.append(f"{app.street_type} {app.street}")
        if app.house_number:
            address_parts.append(f"д. {app.house_number}")
        if app.corpus:
            address_parts.append(f"корп. {app.corpus}")
        if app.apartment_number:
            address_parts.append(f"кв. {app.apartment_number}")
        return ", ".join(address_parts)

    def _fill_docx_template(self, app: ApplicationInfo) -> str:
        doc = Document(self.template_path)

        # Поля для автозамены
        fields = {
            "ИНН": app.inn,
            "Дата рождения": self._format_date(app.birth_date),
            "Телефон": app.phone_number,
            "Электронная почта": getattr(app, "email", ""),
            "Фамилия Имя Отчество": f"{app.surname} {app.name} {app.patronymic}",
            "Адрес регистрации": self._format_address(app),
            "Место работы": getattr(app, "workplace", ""),
            "Должность": getattr(app, "position", ""),
            "Заработная плата": getattr(app, "salary", ""),
            "Адрес работы": getattr(app, "work_address", ""),
            "Образование": getattr(app, "education", ""),
            "Финансовое положение": getattr(app, "financial_status", ""),
            "Семейное положение": getattr(app, "marital_status", ""),
            "Дата трудоустройства": self._format_date(getattr(app, "employment_date", None)),
            "Рабочий телефон": getattr(app, "work_phone", ""),
            "Члены семьи": getattr(app, "family_members", ""),
            "Социальный статус супруга/и": getattr(app, "spouse_status", ""),
            "ФОТО": getattr(app, "photo_path", ""),
            "Данные паспорта": (
                f"{app.type_of_certificate} {app.documents_series} {app.document_number}\n"
                f"Выдан: {app.issued_by}\n"
                f"Дата выдачи: {self._format_date(app.issued_at)}\n"
                f"Страна: {app.country}"
            )
        }

        # Замена в параграфах
        for paragraph in doc.paragraphs:
            for key, val in fields.items():
                if key in paragraph.text:
                    paragraph.text = f"{key}\n{val}"
                    for run in paragraph.runs:
                        run.font.size = Pt(11)

        # Замена в таблицах
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, val in fields.items():
                        if key in cell.text:
                            cell.text = f"{key}\n{val}"

        # Вставляем сегодняшнюю дату в конце документа
        today_str = datetime.today().strftime("%d.%m.%Y")
        for paragraph in doc.paragraphs:
            if paragraph.text.strip() == "Дата":
                paragraph.text = f"Дата {today_str}"
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(11)

        # Сохраняем результат
        timestamp = int(time.time())
        random_str = os.urandom(4).hex()
        filename = sanitize_filename(f"poll_{app.inn}_{timestamp}_{random_str}.docx")
        date_dir = os.path.join(
            self.automation_configs.def_out_paths.docx_reports_file_path,
            get_current_date(self.month, self.year)
        )
        os.makedirs(date_dir, exist_ok=True)
        filepath = os.path.join(date_dir, filename)
        doc.save(filepath)
        return filepath

    def _create_zip_archive(self, generated_files: List[str]) -> str:
        try:
            zip_filename = f"polls_{int(time.time())}.zip"
            zip_path = os.path.join(
                self.automation_configs.def_out_paths.zip_reports_file_path,
                zip_filename
            )
            os.makedirs(os.path.dirname(zip_path), exist_ok=True)
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                added_files = set()
                for file_path in generated_files:
                    if os.path.exists(file_path):
                        arcname = os.path.basename(file_path)
                        if arcname not in added_files:
                            zipf.write(file_path, arcname)
                            added_files.add(arcname)
            for file_path in generated_files:
                if os.path.exists(file_path):
                    os.remove(file_path)
            return zip_path
        except Exception as e:
            logger.error(f"Error creating ZIP archive: {str(e)}")
            raise

    def create_reports_docx(self, application_ids: List[int]) -> str:
        OP = f"{self.OP}.create_reports_docx"
        try:
            applications = self._get_application_data(application_ids)
            if not applications:
                raise NotFoundError("No application data found for given IDs")
            os.makedirs(
                os.path.join(
                    self.automation_configs.def_out_paths.docx_reports_file_path,
                    get_current_date(self.month, self.year)
                ),
                exist_ok=True
            )
            generated_files = []
            for app in applications:
                try:
                    file_path = self._fill_docx_template(app)
                    generated_files.append(file_path)
                    logger.info(f"Generated DOCX: {file_path}")
                except Exception as e:
                    logger.error(f"Error generating DOCX for application {app.id}: {str(e)}")
                    continue
            if not generated_files:
                raise NotFoundError("No DOCX files were generated")
            return self._create_zip_archive(generated_files)
        except Exception as e:
            logger.error(f"[{OP}] Error: {str(e)}")
            raise
